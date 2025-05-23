import streamlit as st
import fitz
from supabase import create_client, Client
from datetime import datetime
from docx import Document
from io import BytesIO
import re

# Supabase config
SUPABASE_URL = "https://syrznbowqhvooxwzikhf.supabase.co"
SUPABASE_KEY = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6InN5cnpuYm93cWh2b294d3ppa2hmIiwicm9sZSI6ImFub24iLCJpYXQiOjE3NDcwNjEzMDUsImV4cCI6MjA2MjYzNzMwNX0.IqeOV-3hynzr2mSN9quFlfkBEaqTKF6LwpL6IlmqYoU"
supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)

# Tabela de seções clínicas
secoes_lab = {
    "Bioquímica": ["glicose", "uréia", "creatinina", "potássio", "sódio", "cálcio", "fósforo"],
    "Hematologia": ["hemoglobina", "hematócrito", "vcm", "hcm", "leucócitos", "plaquetas"],
    "Hormônios": ["tsh", "t4", "t3", "pth"],
    "Vitaminas e Metabolismo Mineral": ["vitamina", "b12", "ácido fólico"],
    "Urina Tipo I": ["ph", "densidade", "hemácias", "leucócitos", "proteína"]
}

ruidos = ["cnpj", "crm", "laboratório", "assinatura", "referência", "nota", "método", "validado"]

def extrair_texto(pdf_file):
    texto = ""
    with fitz.open(stream=pdf_file.read(), filetype="pdf") as doc:
        for page in doc:
            texto += page.get_text()
    return texto

def classificar_exames(texto):
    linhas = texto.lower().splitlines()
    dados = {secao: [] for secao in secoes_lab}
    dados["Outros"] = []
    for linha in linhas:
        l = linha.strip()
        if not l or any(p in l for p in ruidos): continue
        adicionou = False
        for secao, termos in secoes_lab.items():
            if any(t in l for t in termos):
                dados[secao].append(l.capitalize())
                adicionou = True
                break
        if not adicionou and ":" in l:
            dados["Outros"].append(l.capitalize())
    return dados

def limpar_texto(texto):
    texto = texto.encode("utf-8", "ignore").decode("utf-8", "ignore")
    texto = re.sub(r'[^\x20-\x7EÀ-ÿ:\.,\-/()\[\] ]+', '', texto)
    return texto.strip()

def gerar_docx_laboratorial(nome, data, dados):
    doc = Document()
    doc.add_heading(f"Exame Laboratorial", 0)
    doc.add_paragraph(f"Paciente: {nome}")
    doc.add_paragraph(f"Data da coleta: {data}")
    doc.add_paragraph("")
    for secao, itens in dados.items():
        if itens:
            doc.add_heading(secao, level=1)
            for item in itens:
                texto_limpo = limpar_texto(item)
                if texto_limpo:
                    doc.add_paragraph(texto_limpo, style="List Bullet")
    output = BytesIO()
    doc.save(output)
    return output.getvalue()

def gerar_docx_imagem(nome, data, texto):
    doc = Document()
    doc.add_heading(f"Laudo de Imagem", 0)
    doc.add_paragraph(f"Paciente: {nome}")
    doc.add_paragraph(f"Data do exame: {data}")
    doc.add_paragraph("")
    for linha in texto.splitlines():
        texto_limpo = limpar_texto(linha)
        if texto_limpo:
            doc.add_paragraph(texto_limpo)
    output = BytesIO()
    doc.save(output)
    return output.getvalue()

# Interface
st.title("Aplicativo de Laudos – Dr. João Batista Zinato")

with st.form("formulario"):
    tipo = st.radio("Tipo de exame", ["Exame Laboratorial", "Exame de Imagem"])
    nome = st.text_input("Nome do paciente")
    cpf = st.text_input("CPF")
    data = st.date_input("Data do exame")
    arquivo_pdf = st.file_uploader("Envie o PDF do laudo", type="pdf")
    enviar = st.form_submit_button("Processar")

if enviar and arquivo_pdf:
    texto = extrair_texto(arquivo_pdf)
    supabase.table("laudos").insert({
        "nome": nome, "cpf": cpf, "data_nasc": None,
        "data_laudo": datetime.now().isoformat(),
        "conteudo": texto
    }).execute()

    if tipo == "Exame Laboratorial":
        dados = classificar_exames(texto)
        docx_file = gerar_docx_laboratorial(nome, str(data), dados)
    else:
        docx_file = gerar_docx_imagem(nome, str(data), texto)

    st.success("Relatório gerado com sucesso.")
    st.download_button("Baixar relatório", docx_file, file_name="relatorio.docx")
